import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from dataclasses import dataclass


@dataclass
class ParsedDocument:
    source: str
    text: str
    metadata: dict


class DocumentParser:
    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        if path.suffix.lower() == '.pdf':
            return self._parse_pdf(path)
        elif path.suffix.lower() in ('.docx', '.doc'):
            return self._parse_docx(path)
        raise ValueError(f'Unsupported format: {path.suffix}')

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        doc = fitz.open(str(path))
        pages = []
        for page in doc:
            text = page.get_text('text')
            if text.strip():
                pages.append(text)
        full_text = '\n\n'.join(pages)
        return ParsedDocument(
            source=str(path),
            text=full_text,
            metadata={'pages': len(doc), 'filename': path.name, 'type': 'pdf'}
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return ParsedDocument(
            source=str(path),
            text='\n\n'.join(paragraphs),
            metadata={'paragraphs': len(paragraphs), 'filename': path.name, 'type': 'docx'}
        )